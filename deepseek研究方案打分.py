from agentscope.agents import ReActAgentV2, UserAgent
from agentscope.service import ServiceToolkit, execute_python_code
import agentscope
import requests
import json
import networkx as nx
import matplotlib.pyplot as plt
from collections import defaultdict
from PIL import Image, ImageDraw, ImageFont
from io import BytesIO
import os
import time
from agentscope.agents import DialogAgent
from agentscope.message import Msg
from agentscope.pipelines import sequential_pipeline
from agentscope import msghub
import agentscope
from docx import Document
from datetime import datetime
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Optional, Union, Sequence
import agentscope
from agentscope.agents import AgentBase, UserAgent
from agentscope.message import Msg
import init_agents
# 加载模型配置

# 创建三个智能体

init_agents.init()

class StreamingAgent(AgentBase):
    """An agent that speaks streaming messages."""

    def __init__(
        self,
        name: str,
        sys_prompt: str,
        model_config_name: str,
    ) -> None:
        super().__init__(
            name=name,
            sys_prompt=sys_prompt,
            model_config_name=model_config_name,
        )

        self.memory.add(Msg(self.name, self.sys_prompt, "system"))

    def reply(self, x: Optional[Union[Msg, Sequence[Msg]]] = None) -> Msg:
        self.memory.add(x)

        prompt = self.model.format(self.memory.get_memory())

        res = self.model(prompt)


        if res.stream is not None:
            self.speak(res.stream)
        else:

            error_msg = "Received empty response from the model"
            self.speak(error_msg)
            msg_returned = Msg(self.name, error_msg, "assistant")
            self.memory.add(msg_returned)
            return msg_returned

        response_text = res.text.strip() if res.text else "No valid response generated"
        msg_returned = Msg(self.name, response_text, "assistant")

        self.memory.add(msg_returned)

        return msg_returned
deepseek = StreamingAgent(
    "score",
    model_config_name="deepseek_config",
        sys_prompt=(
           f"""
           这是markdown格式的提示词，请你按照提示词完成任务
# 智能助手角色设定
**身份**：资深本科生研究方案评审专家  
**专业领域**：跨学科研究方案评估（侧重理工科与社会科学）  
**核心职能**：对本科生研究方案进行分项量化评分，精准定位扣分点  

---

# 评审执行指令
## 1. 评分维度与权重（总分100分）
### 1.1 选题价值（15分）  
- **评分依据**：  
  ✅ 问题清晰度（4分）  
  ✅ 研究必要性论证（4分）  
  ✅ 目标可衡量性（4分）  
  ✅ 潜在学术/应用价值（3分）  
- **扣分点输出格式**：  
  `[选题价值] 扣分项：问题界定模糊（案例："XX影响因素研究"未明确核心变量）`

### 1.2 文献综述（15分）  
- **评分依据**：  
  ✅ 文献覆盖全面性（5分）  
  ✅ 批判性分析深度（5分）  
  ✅ 理论框架适用性（5分）  
- **扣分点输出格式**：  
  `[文献综述] 扣分项：未指出Smith(2020)实验设计的局限性（需对比方法论差异）`

### 1.3 研究设计（30分）  
- **评分依据**：  
  ✅ 方法匹配研究问题（8分）  
  ✅ 样本选择合理性（7分）  
  ✅ 变量操作化定义（8分）  
  ✅ 数据分析计划可行性（7分）  
- **扣分点输出格式**：  
  `[研究设计] 扣分项：未说明如何控制年龄变量（需补充分层抽样设计）`

### 1.4 可行性（20分）  
- **评分依据**：  
  ✅ 时间规划合理性（6分）  
  ✅ 资源获取路径明确（6分）  
  ✅ 技术能力匹配度（5分）  
  ✅ 伦理合规性（3分）  
- **扣分点输出格式**：  
  `[可行性] 扣分项：电镜设备使用未签署实验室预约证明（需补充凭证）`

### 1.5 创新性（10分）  
- **评分依据**：  
  ✅ 问题/方法新颖度（6分）  
  ✅ 预期成果突破性（4分）  
- **扣分点输出格式**：  
  `[创新性] 扣分项：实验方案完全复现Chen等(2023)流程（需调整参数或增加对照组）`

### 1.6 写作规范（10分）  
- **评分依据**：  
  ✅ 逻辑连贯性（4分）  
  ✅ 学术语言准确性（3分）  
  ✅ 参考文献格式（3分）  
- **扣分输出格式**：  
  `[写作规范] 扣分项：图3未标注数据来源（按APA7需补充版权声明）`

---

## 2. 输出规范
1. **结构化输出**：  
   ```json
   #{{
     "总分": 85,
     "分项评分":{{
       "选题价值": 12,
       "文献综述": 13,
       "研究设计": 25,
       "可行性": 16,
       "创新性": 8,
       "写作规范": 9
     }},
     "扣分明细": [
       "[选题价值] 扣分项：...",
       "[研究设计] 扣分项：..."
     ]
   }}
        """
        )

)

# 通过msghub创建一个聊天室，智能体的消息会广播给所有参与者
with agentscope.msghub(
    participants=[deepseek], 
) as hub:
    # 按顺序发言
    #sequential_pipeline([friday, saturday, judgement],x=None)
    hub.add(deepseek)
    user_agent = UserAgent(name="用户")
    x=user_agent(x=None)
    x=deepseek(x)

doc = Document()
doc.add_heading('对话记录', 0)

tmp=0

for i in deepseek.memory.get_memory():
        # 添加轮次标题
        if i.name =="用户":
            tmp=tmp+1
            round_heading = doc.add_paragraph()

        # 添加对话内容
        p = doc.add_paragraph()
        run = p.add_run(f"{i.name}：")
        run.bold = True    
            # 用户消息样式（蓝色）
        if i.name == "研究总监":
            run.font.color.rgb = RGBColor(0, 0, 150)  # 深蓝色
            p.add_run(i.content)
            # AI消息样式（绿色）
        elif i.name == "批判性思维专家":
            run.font.color.rgb = RGBColor(150,0, 0)
            p.add_run(i.content)
        elif i.name == "文献综述专家":
            run.font.color.rgb = RGBColor(0, 150, 0)
            p.add_run(i.content)
        elif i.name == "研究方法专家":
            run.font.color.rgb = RGBColor(150, 100, 0)
            p.add_run(i.content)
        elif i.name == "数据分析专家":
            run.font.color.rgb = RGBColor(0, 100, 150)
            p.add_run(i.content)
        elif i.name == "学术写作专家":
            run.font.color.rgb = RGBColor(150, 100, 150)
            p.add_run(i.content)
        elif i.name == "用户":
            run.font.color.rgb = RGBColor(100, 100, 200)  # 黑色
            p.add_run(i.content)
        else:
            run.font.color.rgb = RGBColor(20, 100, 200)  
            p.add_run(i.content)   
        # 添加轮次分隔线
# 保存文件
doc.save(f"研究方案打分记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx") 