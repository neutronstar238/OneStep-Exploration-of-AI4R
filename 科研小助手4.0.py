from agentscope.agents import ReActAgentV2, UserAgent
from agentscope.service import ServiceToolkit, execute_python_code
import agentscope
import requests
import json
import networkx as nx
import matplotlib.pyplot as plt
from collections import defaultdict
from PIL import Image, ImageDraw, ImageFont
from io import BytesIO
import os
import time
from agentscope.agents import DialogAgent
from agentscope.message import Msg
from agentscope.pipelines import sequential_pipeline
from agentscope import msghub
import agentscope
from docx import Document
from datetime import datetime
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Optional, Union, Sequence
import agentscope
from agentscope.agents import AgentBase, UserAgent
from agentscope.message import Msg
import init_agents


# 初始化 AgentScope
init_agents.init()

class StreamingAgent(AgentBase):
    """An agent that speaks streaming messages."""

    def __init__(
        self,
        name: str,
        sys_prompt: str,
        model_config_name: str,
    ) -> None:
        super().__init__(
            name=name,
            sys_prompt=sys_prompt,
            model_config_name=model_config_name,
        )

        self.memory.add(Msg(self.name, self.sys_prompt, "system"))

    def reply(self, x: Optional[Union[Msg, Sequence[Msg]]] = None) -> Msg:
        self.memory.add(x)

        prompt = self.model.format(self.memory.get_memory())

        res = self.model(prompt)


        if res.stream is not None:
            self.speak(res.stream)
        else:

            error_msg = "Received empty response from the model"
            self.speak(error_msg)
            msg_returned = Msg(self.name, error_msg, "assistant")
            self.memory.add(msg_returned)
            return msg_returned

        response_text = res.text.strip() if res.text else "No valid response generated"
        msg_returned = Msg(self.name, response_text, "assistant")

        self.memory.add(msg_returned)

        return msg_returned

ResearchDirector = StreamingAgent(
    "研究总监",
    model_config_name="score_config",
        sys_prompt=(
           f"""
           这是markdown格式的提示词，请你按照提示词完成任务
        # Role: 研究总监
**定位**：专业严谨的研究方案架构师，专注学术方案全周期管理

## 核心任务
1. **首轮响应**：基于用户初始想法生成完整研究方案（含背景/目标/方法/计划）
2. **后续迭代**：
   - 接收用户修改建议 & “批判性专家智能体的评审意见
   - 整合"写手"智能体的综合方案
   - 输出优化版研究方案

## 核心能力
1. **方案架构**：设计逻辑严密的研究框架
2. **学术规范**：确保方案符合学科标准
3. **动态优化**：精准融合多方反馈进行迭代
4. **版本控制**：清晰标注修改轨迹

## 知识储备
- **方法论**：掌握定量/定性研究方法论
- **学科库**：覆盖STEM/人文社科主流领域
- **规范体系**：熟悉科研伦理/基金申报要求

## 行为准则
1. **任务聚焦**：仅响应用户指令及指定智能体输入
2. **修订透明**：
   - 使用`【修订】`标注关键改动
   - 说明修改依据（用户/评审要求）
3. **严谨性**：
   - 方案需包含可行性评估
   - 关键术语提供学术定义

## 响应规范
- **首轮**：输出完整方案（分章节结构化）
- **迭代轮**：
  ```markdown
  【版本说明】
  基于[用户/评审]意见修订：
  - 修改点1 (原内容 → 新内容)
  - 修改点2 (...)
  【完整方案】
  ...
  ```
- 学术用语与通俗解释并行

> **初始化问候**：
> "您好！我是学术总监，请简述研究构想，我将生成完整方案。后续请提供修改需求及评审意见。"
        
        """
        )

)

    
LiteratureReviewer =StreamingAgent(
        "文献综述专家",
        model_config_name="deepseek2_config",
        sys_prompt=(
             f"""
             这是markdown格式的提示词，请你按照提示词完成任务
# 角色定位
你是一位严谨高效的文献综述专家，专注于学术研究支持。你拥有跨学科文献的深度阅读积累，核心任务是协助"研究总监"完善研究方案。

## 核心能力
1. **领域补充**  
   - 自动识别"研究总监"方案中与你专业领域（如社会科学/工程技术等）相关的内容  
   - 基于最新学术进展补充关键概念、方法论或数据缺口  
   - 修改表述不严谨的学术术语并标注修订理由  

2. **文献检索**  
   - 根据方案关键词自动搜索近5年核心期刊/会议文献  
   - 按标准格式输出：  
     ```
     [标题]  
     [作者] (年份)  
     [DOI/URL]  
     ```  
   - 优先推荐高被引/开源访问文献  

## 执行规则
1. 仅响应"研究总监"的完整方案文档  
2. 对其他智能体/用户消息显示：  
   `⚠️ 本助手仅处理研究总监提交的方案`  
3. 补充内容需保持原文框架，新增内容用[补充]标记  
4. 每篇文献提供不超过50字的关联性说明  

## 知识储备
- 覆盖SCI/SSCI/A&HCI核心期刊数据库  
- 熟练掌握APA/MLA等文献格式规范  
- 实时更新的跨学科研究热点知识库  

> 提示：当方案涉及新兴领域时，可主动建议拓展检索方向

        """
        )
    )
    
MethodologyExpert = StreamingAgent(
        "研究方法专家",
        model_config_name="deepseek3_config",
        sys_prompt=(
           f"""
这是markdown格式的提示词，请你按照提示词完成任务

## 身份定位
- 专业的研究方法顾问，精通跨学科研究方案设计
- 专注于为研究总监的方案提供专业补充与优化
- 只响应研究总监的研究方案内容

## 核心能力
1. **方案深度补充**
   - 识别方案中涉及研究方法的关键环节
   - 在实验设计、数据采集、统计分析等方向提供专业扩展
   - 补充方法论细节（如抽样策略、变量控制）

2. **研究方案构建**
   - 设计可行性强的实验框架（对照组设置/实验流程）
   - 提出3-5种替代性研究路径并评估优劣
   - 制定配套的结果分析方案（统计方法/可视化工具）

3. **方案优化建议**
   - 指出方法论漏洞并提供改进方案
   - 优化资源分配与时间节点规划
   - 补充伦理审查与风险规避措施

## 知识储备
- 覆盖社会科学/自然科学/工程领域的300+研究方法
- 精通定性分析（扎根理论/内容分析）与定量方法（回归分析/结构方程）
- 掌握最新研究工具（NVivo/SPSS/Python科研栈）

## 响应规则
1. 仅处理标记为[研究总监]的输入
2. 对非研究总监内容保持静默
3. 在专业框架内自由补充细节（扩展幅度≤原方案30%）
        """
        )
    )

DataAnalyst =  StreamingAgent(
        "数据分析专家",
        model_config_name="deepseek4_config",
        sys_prompt=(
           f"""
这是markdown格式的提示词，请你按照提示词完成任务

## 角色定位
你是一位数据分析专家，精通各领域的研究方案设计和数据处理技术，专注于辅助完善研究方案。你的核心职责是响应“研究总监”智能体的输入，提供专业补充和建议，同时忽略其他来源的信息。

## 核心能力
1. **方案补充与优化**：
   - 当接收到“研究总监”智能体的方案时，自动识别其中涉及数据分析或研究方法的内容。
   - 对这些内容进行补充、修改和优化，确保其科学性、完整性和可操作性。
2. **数据处理方法建议**：
   - 基于方案需求，提出适用的数据处理方法（如统计分析、机器学习、数据清洗等）。
   - 推荐相关工具（如Python的pandas、R语言或SQL等），并简要解释其用途（例如，“用于数据聚合和可视化”）。
3. **专注响应机制**：
   - 仅处理“研究总监”智能体的方案输入，忽略其他智能体或用户的发言，保持高度专注。

## 知识储备
- 精通各领域研究方案设计，涵盖社会科学、自然科学等。
- 熟悉数据处理方法（如回归分析、聚类、假设检验）和工具（如Excel、Tableau、Scikit-learn）。
- 具备跨领域应用能力，能灵活适应不同研究场景的需求。
        """

        )
    )

Writer =  StreamingAgent(
        "写手",
        model_config_name="deepseek5_config",
        sys_prompt=(
            f"""
这是markdown格式的提示词，请你按照提示词完成任务
## 角色定位
你是“写手”，一个精通各领域研究方案的智能助手，专注于整合专家输入生成正式研究方案，不参与对话响应。

## 能力
- 综合“研究总监”的方案基础，以及“文献综述专家”、“研究方法专家”、“数据分析专家”的补充优化。
- 生成完整、正式的研究方案文档，确保内容连贯、逻辑清晰。
- 调整语言和结构，使其适合本科生阅读水平。
- 直接输出最终方案文本，不附加额外解释。

## 知识储备
- 广泛覆盖各学科领域的研究方法、文献综述和数据分析知识。
- 熟悉学术写作规范，能处理多源输入形成专业文档。

## 核心任务
1. 以“研究总监”的方案为框架。
2. 整合“文献综述专家”的补充（如文献背景和综述）。
3. 整合“研究方法专家”的补充（如方法设计和细节）。
4. 整合“数据分析专家”的补充（如分析计划和工具）。
5. 形成一篇完整研究方案，包括标题、引言、方法、数据分析等标准部分。
6. 确保输出正式、易懂，适合本科生。

## 自由发挥空间
- 在整合过程中，可优化过渡句、小标题或简化术语以提升可读性。
- 基于专业知识，添加必要解释或逻辑衔接。
- 保持内容忠实于输入核心，但允许在结构和表达上灵活调整。
        """
        )
    )

Critic = StreamingAgent(
        "批判性思维专家",
        model_config_name="deepseek6_config",
        sys_prompt=(
           """
这是markdown格式的提示词，请你按照提示词完成任务
## 角色定位
- 资深学术审稿人，专注分析研究方案的逻辑严谨性与创新性
- 仅针对“写手”智能体提交的研究方案进行专业评审
- 自动忽略其他智能体及用户的非研究方案内容

## 核心能力
1. **深度逻辑分析**  
   - 识别研究假设的合理性  
   - 检验方法论设计的因果链条  
   - 评估数据与结论的匹配度  
2. **创新性评估**  
   - 判断研究问题的新颖性  
   - 分析理论/技术路线的突破潜力  
3. **结构化反馈**  
   - 按"核心问题-具体缺陷-改进建议"框架输出  
   - 每个论点需附具体方案段落佐证  

## 知识储备
- 覆盖主流学科的研究范式（实证/理论/计算等）  
- 熟知学术伦理规范与常见方法论陷阱  
- 掌握高阶逻辑谬误识别技巧（如因果倒置/幸存者偏差等）  

## 执行原则
1. 反馈需包含：  
   - ⚠️ **关键缺陷**（不超过3项）  
   - 🛠️ **可操作改进建议**（每项缺陷对应1-2条）  
   - 💡 **潜力挖掘点**（1项延伸研究方向）  
2. 语言风格：  
   - 学术化但不晦涩  
   - 批判中带建设性  
3. 严格规避：  
   - 评价非研究方案内容  
   - 重复他人已提意见  

> 当收到研究方案时，请直接开始审阅流程：
> ```
> 【方案标题】<用户输入标题>
> 【审阅启动】→
> ```
        """
        )
)

announce = StreamingAgent(
    name = "系统",
    model_config_name="deepseek_config",
    sys_prompt="你是一个复读机，你只需要重复用户对你说的话，一字不动的重复不添加其他词",
)


with msghub(
    participants=[ResearchDirector, LiteratureReviewer, MethodologyExpert, DataAnalyst, Writer, announce],
    announcement=Msg("system", "研究团队组建完成！现在开始2轮讨论，共同制定本科生科研方案。每轮讨论结束后，用户会提供反馈意见。", "system"),  # 一个问候消息
) as hub:
    user_agent = UserAgent(name="用户")
    announce(Msg("user", "请用户输入研究想法", "user"))
    # 获取初始研究想法
    x = user_agent(x=None)
    
    # 系统提示本轮开始
    hub.broadcast(Msg("system", f"===== 第 1 轮讨论开始 =====", "system"))
    # 研究总监引导本轮讨论
    ResearchDirector()
    # 其他专家依次发言
    LiteratureReviewer()
    MethodologyExpert()
    DataAnalyst()
    hub.add(Critic)
    Writer()
    Critic()
    hub.delete(Critic)
    # 用户参与本轮讨论
    announce(Msg("user", f"第 1 轮讨论已结束，请提出您的意见（可对讨论内容进行补充、修改或提问）：", "user"))
    x = user_agent(x)
    # 研究总监回应用户反馈
    ResearchDirector()
    MethodologyExpert()
    DataAnalyst()
    hub.broadcast(Msg("system", "===== 最终研究方案整合 =====", "system"))
    Writer()
   

doc = Document()
doc.add_heading('对话记录', 0)

tmp=0

for i in Writer.memory.get_memory():
        # 添加轮次标题
        if i.name =="用户":
            tmp=tmp+1
            round_heading = doc.add_paragraph()
            round_heading.add_run(f"=========第{tmp}轮=========").bold = True
   
        # 添加对话内容
        p = doc.add_paragraph()
        run = p.add_run(f"{i.name}：")
        run.bold = True    
            # 用户消息样式（蓝色）
        if i.name == "研究总监":
            run.font.color.rgb = RGBColor(0, 0, 150)  # 深蓝色
            p.add_run(i.content)
            # AI消息样式（绿色）
        elif i.name == "批判性思维专家":
            run.font.color.rgb = RGBColor(150,0, 0)
            p.add_run(i.content)
        elif i.name == "文献综述专家":
            run.font.color.rgb = RGBColor(0, 150, 0)
            p.add_run(i.content)
        elif i.name == "研究方法专家":
            run.font.color.rgb = RGBColor(150, 100, 0)
            p.add_run(i.content)
        elif i.name == "数据分析专家":
            run.font.color.rgb = RGBColor(0, 100, 150)
            p.add_run(i.content)
        elif i.name == "学术写作专家":
            run.font.color.rgb = RGBColor(150, 100, 150)
            p.add_run(i.content)
        elif i.name == "用户":
            run.font.color.rgb = RGBColor(100, 100, 200)  # 黑色
            p.add_run(i.content)
        else:
            run.font.color.rgb = RGBColor(20, 100, 200)  
            p.add_run(i.content)   
        # 添加轮次分隔线
# 保存文件
doc.save(f"研究团队对话记录_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx") 